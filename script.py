import openpyxl, os, subprocess, pandas as pd
from docx import Document
from pathlib import Path

EXCEL_NAME = 'Plantilla.xlsx'
COLUMN_TITLES = ['Asignatura', 'Nota', 'Status copy']
COLUMN_WIDTHS = [50, 10, 15]

def get_folders(targetDir, finalDir, docName):
  finalDoc = Document()
  docDir = finalDir + docName + ".docx"
  targetDocs = targetDir
  return finalDoc, docDir, targetDocs

def adjust_column_widths(excel_path):
    workbook = openpyxl.load_workbook(excel_path)
    sheet = workbook.active

    for column in sheet.columns:
        max_length = 0
        column = [cell for cell in column]
        for cell in column:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(cell.value)
            except:
                pass
        adjusted_width = (max_length + 2)
        sheet.column_dimensions[column[0].column_letter].width = adjusted_width

    workbook.save(excel_path)
    workbook.close()

def create_template():
    template = openpyxl.Workbook()
    sheet = template.active
    sheet.title = 'Sheet1'
    
    for i, title in enumerate(COLUMN_TITLES):
        cell = chr(65 + i) + '1'
        sheet[cell] = title
        sheet.column_dimensions[chr(65 + i)].width = COLUMN_WIDTHS[i]

    template.save(EXCEL_NAME)
    template.close()
    
    open_excel(EXCEL_NAME)

def open_excel(excelName):
  if os.name == 'nt':  
    subprocess.Popen(['start', 'excel', excelName], shell=True)
  if os.name == 'posix':  
    subprocess.Popen(['xdg-open', excelName])

def copy_paragraphs(source_paragraphs, target_document):
    for source_paragraph in source_paragraphs:
        copy_paragraph(source_paragraph, target_document)

def copy_paragraph(source_paragraph, target_document):
    new_paragraph = target_document.add_paragraph()
    for run in source_paragraph.runs:
        copy_run(run, new_paragraph)

def copy_run(run, new_paragraph):
    new_run = new_paragraph.add_run(run.text)
    new_run.font.size = run.font.size
    new_run.bold = run.bold
    new_run.italic = run.italic
    new_run.underline = run.underline

def copy_tables(source_tables, target_document):
    for source_table in source_tables:
        copy_table(source_table, target_document)

def copy_table(source_table, target_document):
    new_table = target_document.add_table(rows=len(source_table.rows), cols=len(source_table.columns))
    for i, row in enumerate(source_table.rows):
        for j, cell in enumerate(row.cells):
            new_table.cell(i, j).text = cell.text

def process_document(doc_path, finalDoc, df, index):
    doc = Document(doc_path)
    copy_paragraphs(doc.paragraphs, finalDoc)
    copy_tables(doc.tables, finalDoc)
    finalDoc.add_page_break() # Insertar un salto de página después de cada documento
    df.at[index, 'Status copy'] = 'OK'

def run_script(targetDir, finalDir, docName, excelDir):
    finalDoc, docDir, targetDocs = get_folders(targetDir, finalDir, docName)
    root = Path(targetDocs)

    # Crear un diccionario que especifique el tipo de datos para cada columna
    column_dtypes = {'Asignatura': str, 'Nota': float, 'Status copy': str}

    df = pd.read_excel(excelDir, sheet_name='Sheet1', header=0, dtype=column_dtypes)
    for column in df.columns:
        df[column] = df[column].map(lambda x: x.strip() if isinstance(x, str) else x)

    # Cambiar el tipo de datos de la columna 'Status copy' a cadena de caracteres
    df['Status copy'] = df['Status copy'].astype(str)

    # Iterar sobre las filas del DataFrame
    for index, row in df.iterrows():
        asignatura = row['Asignatura']
        nota = row['Nota']
        archivo_encontrado = False
      
        # Buscar archivos en la carpeta que coincidan con el nombre de asignatura y tengan nota >= 3
        for doc_path in root.rglob(f'{asignatura}*.docx'):
            archivo_encontrado = True
            if nota >= 3:
                process_document(doc_path, finalDoc, df, index)
            else:
                df.at[index, 'Status copy'] = 'N/A'

        if nota >= 3 and not archivo_encontrado:
            finalDoc.add_paragraph(f'Falta archivo: {asignatura}')
            df.at[index, 'Status copy'] = 'Falta archivo'

    # Guardar el documento final
    finalDoc.save(docDir)

    # Guardar el status de la copia en el archivo Excel
    df.to_excel(excelDir, index=False)
    adjust_column_widths(excelDir)
    open_excel(excelDir)


targetDir = 'C:/Users/user/Desktop/Trabajos/scriptWord/DocumentosObjetivos'
finalDir = 'C:/Users/user/Desktop/Trabajos/scriptWord/DocumentoFusionado/'
docName = 'Ana Maria'
excelDir = 'C:/Users/user/Desktop/Trabajos/scriptWord/Plantilla.xlsx'

#create_template()
run_script(targetDir, finalDir, docName, excelDir)